import os
import json
import random
import tempfile
import requests
from typing import Tuple, Dict, Any
import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# File parsing libs
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx2txt
except Exception:
    docx2txt = None

# For docx generation
try:
    from docx import Document
except Exception:
    Document = None

# --------------------------
# CONFIG & CONSTANTS
# --------------------------
API_KEY = os.getenv("OPENAI_API_KEY")  # Load from environment variable
API_BASE = "https://api.chatanywhere.tech"
DEFAULT_MODEL = "gpt-3.5-turbo"

AVAILABLE_MODELS = {
    "gpt-3.5-turbo": {"name": "GPT-3.5", "description": "Fast & efficient analysis"},
    "gpt-4o-mini": {"name": "GPT-4o", "description": "More detailed analysis"},
    "gpt-5-mini": {"name": "GPT-5", "description": "Best results, limited usage"},
    "deepseek-v3": {"name": "DeepSeek", "description": "Specialized for resumes"}
}

# ATS Analysis Categories
ATS_CATEGORIES = {
    "Content Quality": [
        "Achievement Focus",
        "Action Verbs",
        "Quantified Results",
        "Industry Keywords",
        "Role-specific Skills"
    ],
    "Format & Structure": [
        "ATS Readability",
        "Section Headers",
        "Bullet Points",
        "Consistent Formatting",
        "Contact Info"
    ],
    "Experience Match": [
        "Job Alignment",
        "Skills Coverage",
        "Required Qualifications",
        "Experience Level",
        "Domain Knowledge"
    ]
}


# --------------------------
# Helper: Chat Completion
# --------------------------
def make_chat_completion(messages, model=DEFAULT_MODEL, temperature=0.0, max_tokens=1000) -> Dict[str, Any]:
    """
    Generic function to call the ChatAnywhere (OpenAI-style) chat completions endpoint.
    Returns the parsed JSON response.
    """
    if not API_KEY:
        raise RuntimeError("API_KEY not set. Set OPENAI_API_KEY env var or API_KEY variable in code.")
    url = f"{API_BASE.rstrip('/')}/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=60)
    try:
        resp.raise_for_status()
    except Exception as e:
        # Surface helpful message
        raise RuntimeError(f"API call failed: {e}\nResponse: {resp.text}")
    return resp.json()


# --------------------------
# File text extraction
# --------------------------
def extract_text_from_file(uploaded_file) -> str:
    """
    Accepts a streamlit uploaded file-like object
    Supports PDF, DOCX, TXT.
    Returns plain text string.
    """
    fname = uploaded_file.name.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp:
        tmp.write(uploaded_file.getbuffer())
        tmp.flush()
        tmp_path = tmp.name

    text = ""
    try:
        if fname.endswith(".pdf"):
            if pdfplumber is None:
                raise RuntimeError("pdfplumber not installed. Add pdfplumber in requirements.")
            with pdfplumber.open(tmp_path) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages)
        elif fname.endswith(".docx") or fname.endswith(".doc"):
            if docx2txt is None:
                raise RuntimeError("docx2txt not installed. Add docx2txt in requirements.")
            text = docx2txt.process(tmp_path) or ""
        else:
            # txt or fallback
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
    return text.strip()


# --------------------------
# Text chunking helpers for large documents
# --------------------------
def estimate_tokens(text: str) -> int:
    """Estimate token count (rough approximation: 1 token ≈ 4 characters)"""
    return len(text) // 4


def chunk_text(text: str, max_tokens: int=3000) -> list:
    """Split text into smaller chunks that fit within token limits"""
    if estimate_tokens(text) <= max_tokens:
        return [text]
    
    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_tokens = 0
    
    for line in lines:
        line_tokens = estimate_tokens(line)
        
        if current_tokens + line_tokens > max_tokens and current_chunk:
            # Save current chunk and start new one
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_tokens = line_tokens
        else:
            current_chunk.append(line)
            current_tokens += line_tokens
    
    # Add the last chunk
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks


def chunk_resume_and_jd(resume_text: str, jd_text: str, max_tokens: int=2800) -> list:
    """Create chunks that include both resume and JD portions while staying under token limit"""
    jd_tokens = estimate_tokens(jd_text)
    resume_tokens = estimate_tokens(resume_text)
    
    # If both fit together, return as single chunk
    if jd_tokens + resume_tokens <= max_tokens:
        return [(resume_text, jd_text)]
    
    # If JD is too long, truncate it
    if jd_tokens > max_tokens // 2:
        jd_text = jd_text[:max_tokens * 2]  # Rough truncation
        jd_tokens = estimate_tokens(jd_text)
    
    # Chunk the resume
    available_tokens = max_tokens - jd_tokens - 100  # Leave buffer
    resume_chunks = chunk_text(resume_text, available_tokens)
    
    # Return list of (resume_chunk, jd_text) tuples
    return [(chunk, jd_text) for chunk in resume_chunks]


# --------------------------
# Prompting helpers
# --------------------------
SCORE_PROMPT = """You are an expert career coach and resume reviewer with experience in recruiting and ATS.
Given the user's resume text and the job description, produce a JSON object with these fields:
- score: integer 0-100 (how good the resume matches the JD and ATS readiness)
- strengths: list of short strings (3-6)
- weaknesses: list of short strings (3-8)
- improvement_tips: list of actionable suggestions (5 items) (formatting, keywords, achievements)
- missing_keywords: list of important keywords/phrases present in JD but missing in resume (up to 12)
- overall_summary: one short paragraph (max 80 words)

Input JSON:
{{
  "resume_text": "{resume_text}",
  "job_description": "{job_description}"
}}

Return only valid JSON (no commentary). If you cannot match fields, still return a JSON with empty lists and a score.
"""

TAILOR_PROMPT = """You are an expert resume writer. Using the user's resume and the job description, produce:
1) A tailored resume in plain text that:
   - Reorders sections for maximum impact
   - Rewrites bullet points to emphasize achievements and use JD keywords
   - Keeps length roughly the same (don't invent jobs)
   - Uses result-driven bullets (use metrics where present)
Return result under a JSON with keys:
{
  "tailored_resume": "<full resume as plain text>"
}
Do not add anything else outside the JSON.
"""

COVER_PROMPT = """You are an expert career coach. Write a short cover-note (3-5 sentences) tailored to the job description and the candidate's resume emphasizing fit and one key achievement. Return as plain text only.
"""

CHUNK_SCORE_PROMPT = """You are an expert career coach analyzing a PART of a resume against a job description.
This is chunk {chunk_num} of {total_chunks} from the resume.

Resume section:
{resume_chunk}

Job Description:
{job_description}

Analyze this section and return JSON with:
- partial_score: integer 0-100 for this section's match quality
- section_strengths: list of strengths found in this section (2-4 items)
- section_weaknesses: list of weaknesses in this section (2-4 items)
- section_keywords: list of keywords from JD found in this section
- missing_keywords: list of important JD keywords missing from this section
- section_summary: brief summary of this section's content and quality

Return only valid JSON.
"""

COMBINE_CHUNKS_PROMPT = """You are combining analysis results from multiple resume sections. 
Here are the individual chunk analyses:

{chunk_results}

Job Description:
{job_description}

Combine these into a final comprehensive analysis with:
- score: integer 0-100 overall resume score
- strengths: list of overall resume strengths (4-6 items)
- weaknesses: list of overall weaknesses (4-6 items)  
- improvement_tips: list of actionable suggestions (5 items)
- missing_keywords: list of important missing keywords (up to 12)
- overall_summary: comprehensive summary paragraph

Return only valid JSON.
"""

COMPARISON_PROMPT = """You are an expert resume reviewer. Compare these two resumes and provide a detailed analysis:

Resume A:
{resume_a}

Resume B:
{resume_b}

Job Description (if provided):
{job_description}

Provide a JSON response with:
- overall_comparison: string summary of which is stronger and why
- resume_a_strengths: list of Resume A's advantages
- resume_b_strengths: list of Resume B's advantages  
- resume_a_score: integer 0-100 overall quality score
- resume_b_score: integer 0-100 overall quality score
- recommendations: list of specific improvements for each resume
- winner: "A" or "B" or "Tie" based on overall assessment

Return only valid JSON.
"""


# --------------------------
# ATS Analysis Functions
# --------------------------
def create_ats_analysis(score_data):
    """Main function to handle all ATS analysis"""
    # Get base score
    base_score = score_data.get("score", 0)
    
    # Calculate category scores
    category_scores = {}
    for category in ATS_CATEGORIES.keys():
        adjustment = len(score_data.get("strengths", [])) - len(score_data.get("weaknesses", []))
        adjustment *= 5
        category_scores[category] = min(100, max(0, base_score + adjustment))
    
    # Calculate detailed scores
    for category, subcats in ATS_CATEGORIES.items():
        category_base = category_scores[category]
        for subcat in subcats:
            score_key = f"{category.lower()}_{subcat.lower().replace(' ', '_')}"
            variation = random.randint(-5, 5)
            score_data[score_key] = min(100, max(0, category_base + variation))
    
    return score_data


def create_ats_visualizations(score_data):
    """Create all ATS-related visualizations"""
    # Create radar chart
    radar_fig = go.Figure()
    categories = []
    scores = []
    
    for category, subcats in ATS_CATEGORIES.items():
        for subcat in subcats:
            label = f"{category}:<br>{subcat}"
            categories.append(label)
            score_key = f"{category.lower()}_{subcat.lower().replace(' ', '_')}"
            score = score_data.get(score_key, score_data.get("score", 50))
            scores.append(min(100, max(0, score)))

    radar_fig.add_trace(go.Scatterpolar(
        r=scores,
        theta=categories,
        fill='toself',
        name='Resume Score',
        line=dict(color='#2E86C1', width=2),
        fillcolor='rgba(46, 134, 193, 0.3)'
    ))

    radar_fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                tickfont=dict(size=10),
                gridcolor='rgba(0,0,0,0.1)',
                ticksuffix="%"
            ),
            angularaxis=dict(
                tickfont=dict(size=8),
                rotation=90,
                direction="clockwise"
            )
        ),
        showlegend=False,
        title="ATS Score Breakdown",
        height=600,
        margin=dict(t=100, b=50)
    )
    
    # Create skills match chart
    missing_keywords = score_data.get("missing_keywords", [])
    matched_keywords = score_data.get("matched_keywords", score_data.get("strengths", []))
    
    match_fig = None
    if missing_keywords or matched_keywords:
        total = len(missing_keywords) + len(matched_keywords)
        match_percent = (len(matched_keywords) / total * 100) if total > 0 else 0
        
        match_fig = go.Figure()
        match_fig.add_trace(go.Bar(
            x=["Skills Match"],
            y=[match_percent],
            name="Matched Keywords",
            marker_color="rgb(46, 134, 193)",
            hovertemplate="Matched: %{y:.1f}%<br>Keywords: " + 
                         ", ".join(matched_keywords[:5]) + 
                         ("..." if len(matched_keywords) > 5 else "")
        ))
        match_fig.add_trace(go.Bar(
            x=["Skills Match"],
            y=[100 - match_percent],
            name="Missing Keywords",
            marker_color="rgba(231, 76, 60, 0.7)",
            hovertemplate="Missing: %{y:.1f}%<br>Keywords: " + 
                         ", ".join(missing_keywords[:5]) + 
                         ("..." if len(missing_keywords) > 5 else "")
        ))
        
        match_fig.update_layout(
            barmode='stack',
            title="Skills Match Analysis",
            yaxis_title="Match Percentage",
            showlegend=True,
            height=300,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
    
    return radar_fig, match_fig


def create_comparison_visualization(comparison_data):
    """Create comparison visualization between two resumes"""
    scores = [
        comparison_data.get("resume_a_score", 50),
        comparison_data.get("resume_b_score", 50)
    ]
    
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=["Resume A", "Resume B"],
        y=scores,
        marker_color=['#3498db', '#e74c3c'],
        text=[f"{score}%" for score in scores],
        textposition='auto',
    ))
    
    fig.update_layout(
        title="Resume Comparison Scores",
        yaxis_title="Overall Score",
        yaxis=dict(range=[0, 100]),
        height=400
    )
    
    return fig


# --------------------------
# API Call Functions
# --------------------------
def call_for_score(resume_text: str, jd: str, model=DEFAULT_MODEL):
    """Score resume with chunking support for large documents"""
    
    # Check if we need to chunk
    total_tokens = estimate_tokens(resume_text + jd)
    
    if total_tokens <= 3500:  # Safe threshold for free API
        # Use original single-call method
        return call_for_score_single(resume_text, jd, model)
    else:
        # Use chunked analysis
        return call_for_score_chunked(resume_text, jd, model)


def call_for_score_single(resume_text: str, jd: str, model=DEFAULT_MODEL):
    """Original single-call scoring method"""
    prompt = SCORE_PROMPT.format(resume_text=resume_text.replace("\n", "\\n"), job_description=jd.replace("\n", "\\n"))
    messages = [
        {"role": "system", "content": "You are an assistant that strictly returns JSON when asked."},
        {"role": "user", "content": prompt}
    ]
    resp = make_chat_completion(messages, model=model, temperature=0.0, max_tokens=800)
    return parse_json_response(resp)


def call_for_score_chunked(resume_text: str, jd: str, model=DEFAULT_MODEL):
    """Chunked scoring method for large resumes"""
    
    # Create chunks
    chunks = chunk_resume_and_jd(resume_text, jd, max_tokens=2800)
    chunk_results = []
    
    # Analyze each chunk
    for i, (resume_chunk, jd_chunk) in enumerate(chunks, 1):
        prompt = CHUNK_SCORE_PROMPT.format(
            chunk_num=i,
            total_chunks=len(chunks),
            resume_chunk=resume_chunk.replace("\n", "\\n"),
            job_description=jd_chunk.replace("\n", "\\n")
        )
        
        messages = [
            {"role": "system", "content": "You are an assistant that strictly returns JSON when asked."},
            {"role": "user", "content": prompt}
        ]
        
        resp = make_chat_completion(messages, model=model, temperature=0.0, max_tokens=600)
        chunk_result = parse_json_response(resp)
        chunk_results.append(chunk_result)
    
    # Combine results
    combined_prompt = COMBINE_CHUNKS_PROMPT.format(
        chunk_results=json.dumps(chunk_results, indent=2),
        job_description=jd.replace("\n", "\\n")
    )
    
    messages = [
        {"role": "system", "content": "You are an assistant that combines analysis results and returns JSON."},
        {"role": "user", "content": combined_prompt}
    ]
    
    resp = make_chat_completion(messages, model=model, temperature=0.0, max_tokens=1000)
    return parse_json_response(resp)


def parse_json_response(resp):
    """Helper function to parse JSON from API response"""
    try:
        assistant_text = resp["choices"][0]["message"]["content"]
        # Attempt to find JSON substring
        start = assistant_text.find("{")
        end = assistant_text.rfind("}") + 1
        if start != -1 and end != -1:
            json_text = assistant_text[start:end]
            data = json.loads(json_text)
            return data
        else:
            # As fallback, try to parse whole message
            return json.loads(assistant_text)
    except Exception as e:
        raise RuntimeError(f"Failed to parse JSON response: {e}\nRaw response: {resp}")


def call_for_tailor(resume_text: str, jd: str, model=DEFAULT_MODEL):
    """Tailor resume with chunking support for large documents"""
    
    # Check if we need to chunk
    total_tokens = estimate_tokens(resume_text + jd)
    
    if total_tokens <= 3500:  # Safe threshold for free API
        # Use original method for smaller resumes
        messages = [
            {"role": "system", "content": "You are an expert resume writer. Output a JSON object with a 'tailored_resume' key."},
            {"role": "user", "content": f"Resume:\n{resume_text}\n\nJob Description:\n{jd}"}
        ]
        resp = make_chat_completion(messages, model=model, temperature=0.2, max_tokens=1500)
        return parse_json_response(resp)
    else:
        # For large resumes, provide a simplified tailoring approach
        # Truncate resume to fit within limits while preserving key sections
        resume_lines = resume_text.split('\n')
        
        # Keep important sections (contact info, summary, experience)
        important_lines = []
        section_keywords = ['contact', 'summary', 'experience', 'education', 'skills']
        
        current_section = None
        for line in resume_lines:
            line_lower = line.lower().strip()
            
            # Check if this is a section header
            if any(keyword in line_lower for keyword in section_keywords):
                current_section = line
                important_lines.append(line)
            elif current_section and len('\n'.join(important_lines)) < 2000:  # Keep within reasonable size
                important_lines.append(line)
        
        condensed_resume = '\n'.join(important_lines)
        
        # If still too long, truncate further
        if estimate_tokens(condensed_resume + jd) > 3500:
            condensed_resume = condensed_resume[:2000]
        
        messages = [
            {"role": "system", "content": "You are an expert resume writer. Output a JSON object with a 'tailored_resume' key. Focus on the most important sections provided."},
            {"role": "user", "content": f"Resume (key sections):\n{condensed_resume}\n\nJob Description:\n{jd}"}
        ]
        resp = make_chat_completion(messages, model=model, temperature=0.2, max_tokens=1200)
        return parse_json_response(resp)


def call_for_cover(resume_text: str, jd: str, model=DEFAULT_MODEL):
    """Generate cover note with text truncation for large inputs"""
    
    # Truncate inputs if too long
    if estimate_tokens(resume_text + jd) > 3000:
        resume_text = resume_text[:1500]  # Keep first part of resume
        jd = jd[:1000]  # Truncate job description if needed
    
    messages = [
        {"role": "system", "content": "You are an expert career coach."},
        {"role": "user", "content": f"Write a short cover note based on this resume and job description.\n\nResume:\n{resume_text}\n\nJD:\n{jd}"}
    ]
    resp = make_chat_completion(messages, model=model, temperature=0.3, max_tokens=200)
    return resp["choices"][0]["message"]["content"]


def call_for_comparison(resume_a: str, resume_b: str, jd: str="", model=DEFAULT_MODEL):
    """Compare two resumes using AI with text truncation for large inputs"""
    
    # Estimate total tokens and truncate if necessary
    total_tokens = estimate_tokens(resume_a + resume_b + jd)
    
    if total_tokens > 3500:
        # Truncate resumes proportionally
        max_resume_length = 1200  # Characters per resume
        resume_a = resume_a[:max_resume_length] + "..." if len(resume_a) > max_resume_length else resume_a
        resume_b = resume_b[:max_resume_length] + "..." if len(resume_b) > max_resume_length else resume_b
        
        # Truncate job description if still too long
        if jd and len(jd) > 800:
            jd = jd[:800] + "..."
    
    prompt = COMPARISON_PROMPT.format(
        resume_a=resume_a.replace("\n", "\\n"),
        resume_b=resume_b.replace("\n", "\\n"),
        job_description=jd.replace("\n", "\\n") if jd else "Not provided"
    )
    messages = [
        {"role": "system", "content": "You are an expert resume reviewer that returns JSON responses."},
        {"role": "user", "content": prompt}
    ]
    resp = make_chat_completion(messages, model=model, temperature=0.1, max_tokens=1200)
    return parse_json_response(resp)


def ensure_string(data):
    """Ensure data is a string for download"""
    if isinstance(data, dict):
        return json.dumps(data, indent=2)
    return str(data)


def format_resume_output(resume_text: str) -> str:
    """Format resume text with proper section detection and spacing"""
    sections = []
    current_section = []
    
    # Common section headers in resumes
    SECTION_KEYWORDS = {
        "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS",
        "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS",
        "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
        "TECHNICAL SKILLS", "PUBLICATIONS", "LANGUAGES"
    }
    
    for line in resume_text.splitlines():
        line = line.strip()
        # Enhanced section header detection
        is_header = (
            line.upper() in SECTION_KEYWORDS or
            any(line.upper().startswith(k) for k in SECTION_KEYWORDS) or
            line.endswith(':') or
            (line.isupper() and len(line) > 3)  # Avoid short uppercase words
        )
        
        if is_header and current_section:
            # Complete previous section
            sections.append('\n'.join(current_section))
            current_section = []
            # Start new section with formatted header
            current_section.append(f"\n{line.upper()}")
        elif line:
            # Format bullet points consistently
            if line.lstrip().startswith(('•', '-', '*', '○', '→')):
                line = '• ' + line.lstrip('•-*○→ \t')
            current_section.append(line)
    
    # Add final section
    if current_section:
        sections.append('\n'.join(current_section))
    
    # Join sections with proper spacing
    formatted_text = '\n\n'.join(sections)
    # Clean up multiple blank lines
    formatted_text = '\n'.join(line for line in formatted_text.splitlines() if line.strip() or line == '')
    while '\n\n\n' in formatted_text:
        formatted_text = formatted_text.replace('\n\n\n', '\n\n')
    
    return formatted_text.strip()


# --------------------------
# PAGE FUNCTIONS
# --------------------------
def resume_analysis_page():
    """Main resume analysis and tailoring page"""
    st.title("🔍 Resume Score & Tailor")
    st.markdown("Upload your resume and paste the job description (JD). This app will score your resume, give improvements, and produce a tailored resume using AI.")

    # Model Selection for this page
    with st.sidebar:
        st.markdown("### 🤖 Model Selection")
        model_choice = st.selectbox(
            "Choose Analysis Model",
            options=list(AVAILABLE_MODELS.keys()),
            format_func=lambda x: AVAILABLE_MODELS[x]["name"],
            index=list(AVAILABLE_MODELS.keys()).index(DEFAULT_MODEL),
            help="Select the AI model to use for analysis. Each model has different capabilities."
        )
        st.caption(f"💡 {AVAILABLE_MODELS[model_choice]['description']}")

    # Upload resume and JD
    st.subheader("1) Upload resume")
    uploaded = st.file_uploader("Upload PDF / DOCX / TXT resume", type=["pdf", "docx", "txt"])
    st.subheader("2) Paste Job Description")
    jd_text = st.text_area("Paste the Job Description (JD) here", height=220)

    col1, col2 = st.columns(2)
    with col1:
        analyze_btn = st.button("🔎 Analyze & Score")
    with col2:
        tailor_btn = st.button("✍️ Generate Tailored Resume")

    # Process user actions
    if analyze_btn:
        if not uploaded:
            st.warning("Please upload a resume file first.")
        elif not jd_text or len(jd_text.strip()) < 20:
            st.warning("Please paste a detailed job description (at least ~20 characters).")
        else:
            with st.spinner("Extracting resume text..."):
                try:
                    resume_text = extract_text_from_file(uploaded)
                except Exception as e:
                    st.error(f"Failed to extract text from file: {e}")
                    resume_text = None

            if resume_text:
                # Check if we need to chunk and inform user
                total_tokens = estimate_tokens(resume_text + jd_text)
                if total_tokens > 3500:
                    st.info("📄 Large resume detected! Processing in smaller chunks to ensure compatibility with the free API...")
                    st.caption("💡 This may take a few extra moments but ensures complete analysis.")
                else:
                    st.info("Contacting AI API to score the resume...")
                
                try:
                    score_data = call_for_score(resume_text, jd_text, model_choice)
                except Exception as e:
                    st.error(f"Error from API: {e}")
                    if "token" in str(e).lower() or "4096" in str(e):
                        st.error("💡 **Tip**: Your resume might be too long. Try a shorter resume or remove some sections temporarily.")
                    score_data = None

                if score_data:
                    score_data = create_ats_analysis(score_data)
                    radar_chart, skills_chart = create_ats_visualizations(score_data)
                    
                    # Display results
                    score_col, viz_col = st.columns([1, 2])
                    with score_col:
                        st.subheader("📊 Resume Score")
                        score = score_data.get("score", "N/A")
                        st.metric("Match Score (0-100)", score)
                        
                        # Add color-coded score interpretation
                        if isinstance(score, (int, float)):
                            if score >= 80:
                                st.success("Strong Match! 🌟")
                            elif score >= 60:
                                st.info("Good Match 👍")
                            else:
                                st.warning("Needs Improvement 🔨")
                    with viz_col:
                        st.plotly_chart(radar_chart, use_container_width=True)
                    
                    if skills_chart:
                        st.plotly_chart(skills_chart, use_container_width=True)
                    
                    # Continue with existing sections
                    st.subheader("Strengths")
                    for s in score_data.get("strengths", []):
                        st.success(f"• {s}")
                    st.subheader("Weaknesses")
                    for w in score_data.get("weaknesses", []):
                        st.error(f"• {w}")
                    st.subheader("Improvement Tips")
                    for t in score_data.get("improvement_tips", []):
                        st.write(f"• {t}")
                    st.subheader("Missing Keywords (from JD)")
                    if score_data.get("missing_keywords"):
                        st.write(", ".join(score_data.get("missing_keywords")))
                    else:
                        st.write("None detected")
                    st.subheader("Summary")
                    st.write(score_data.get("overall_summary", ""))
                    # Save for later tailoring
                    st.session_state["last_resume_text"] = resume_text
                    st.session_state["last_jd"] = jd_text

    if tailor_btn:
        # ensure we have resume_text
        if "last_resume_text" not in st.session_state:
            if not uploaded:
                st.warning("Please upload a resume and click Analyze first (or paste JD and click Analyze).")
            else:
                with st.spinner("Extracting resume text..."):
                    try:
                        resume_text = extract_text_from_file(uploaded)
                        st.session_state["last_resume_text"] = resume_text
                        st.session_state["last_jd"] = jd_text
                    except Exception as e:
                        st.error(f"Failed to extract text from file: {e}")
        if "last_resume_text" in st.session_state:
            resume_text = st.session_state["last_resume_text"]
            jd_text_local = st.session_state.get("last_jd", jd_text)
            
            # Check if we need to process a large resume
            total_tokens = estimate_tokens(resume_text + jd_text_local)
            if total_tokens > 3500:
                st.info("📝 Processing large resume for tailoring. Focusing on key sections...")
            else:
                st.info("Asking AI to generate a tailored resume...")
                
            try:
                tailored_json = call_for_tailor(resume_text, jd_text_local, model_choice)
                tailored_resume = tailored_json.get("tailored_resume", "")
            except Exception as e:
                st.error(f"Error generating tailored resume: {e}")
                if "token" in str(e).lower() or "4096" in str(e):
                    st.error("💡 **Tip**: Resume too long for tailoring. Try analyzing first to save a shorter version.")
                tailored_resume = ""

            if tailored_resume:
                st.subheader("✏️ Tailored Resume (Preview)")
                
                # Format the resume text
                tailored_text = ensure_string(tailored_resume)
                formatted_text = format_resume_output(tailored_text)
                
                # Create columns for better layout
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.text_area("Preview", formatted_text, height=500, key="tailored_preview")
                
                with col2:
                    st.markdown("### Download Options")
                    st.download_button(
                        "📄 Download as TXT",
                        formatted_text,
                        file_name="tailored_resume.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    if Document is not None:
                        # Create formatted docx
                        doc = Document()
                        for section in formatted_text.split('\n\n'):
                            if section.strip():
                                if section.isupper() or section.endswith(':'):
                                    doc.add_heading(section.strip(':'), level=2)
                                else:
                                    doc.add_paragraph(section)
                    
                        tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        doc.save(tmp_docx.name)
                        tmp_docx.flush()
                        
                        with open(tmp_docx.name, "rb") as f:
                            st.download_button(
                                "📑 Download as DOCX",
                                f.read(),
                                file_name="tailored_resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        try:
                            os.unlink(tmp_docx.name)
                        except:
                            pass

                # Enhanced cover note section
                st.subheader("📝 Short Cover Note")
                try:
                    cover = call_for_cover(resume_text, jd_text_local, model_choice)
                    cover_text = ensure_string(cover).strip()
                    if cover_text and len(cover_text) > 10:  # Basic validation
                        st.markdown(
                            f"""
                            <div style='padding: 1rem; border-left: 4px solid #2E86C1; background-color: #f0f2f6;'>
                            {cover_text}
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
                    else:
                        st.warning("Could not generate a meaningful cover note. Please try again.")
                except Exception as e:
                    st.error(f"Error generating cover note: {str(e)}")

    # Footer: quick tips
    st.markdown("---")
    st.markdown("**Tips:**\n- For best results, paste the full job description (responsibilities + required skills). \n- If the model returns malformed JSON, try a shorter resume or split large files. \n- This app depends on the third-party API endpoint — if requests fail, double-check the API_BASE and your API key.")


def skill_gap_analysis_page():
    """Skill Gap Analysis page - Shows missing skills vs. job requirements"""
    
    st.title("🎯 Skill Gap Analysis")
    st.markdown("**Identify missing skills and get targeted recommendations for improvement**")
    
    # Model Selection for this page
    with st.sidebar:
        st.subheader("⚙️ Analysis Settings")
        selected_model = st.selectbox(
            "Select AI Model:",
            options=list(AVAILABLE_MODELS.keys()),
            format_func=lambda x: f"{AVAILABLE_MODELS[x]['name']} - {AVAILABLE_MODELS[x]['description']}",
            key="skill_gap_model"
        )
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Your Resume")
        resume_file = st.file_uploader(
            "Upload your resume",
            type=['pdf', 'docx', 'txt'],
            key="skill_gap_resume"
        )
        
        if resume_file:
            try:
                resume_text = extract_text_from_file(resume_file)
                st.success(f"Resume loaded: {len(resume_text)} characters")
                with st.expander("Preview Resume Text"):
                    st.text_area("Resume Content:", resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text, height=200, disabled=True)
            except Exception as e:
                st.error(f"Error reading resume: {e}")
                resume_text = None
        else:
            resume_text = None
    
    with col2:
        st.subheader("💼 Job Requirements")
        job_description = st.text_area(
            "Paste the job description or requirements:",
            height=200,
            placeholder="Enter the complete job description including required skills, qualifications, and experience..."
        )
        
        # Optional: Company name for better context
        company_name = st.text_input("Company Name (optional):", placeholder="e.g., Google, Microsoft")
    
    if st.button("🔍 Analyze Skill Gaps", type="primary"):
        if not resume_text:
            st.error("Please upload your resume first!")
        elif not job_description.strip():
            st.error("Please provide the job description!")
        else:
            with st.spinner("Analyzing skill gaps..."):
                try:
                    # Prepare the analysis prompt
                    company_context = f" at {company_name}" if company_name else ""
                    messages = [
                        {
                            "role": "system",
                            "content": """You are an expert career coach and skill gap analyst. Analyze the candidate's resume against job requirements to identify skill gaps and provide actionable recommendations.

Provide your analysis in the following JSON format:
{
    "skill_gaps": {
        "critical_missing": ["skill1", "skill2", ...],
        "nice_to_have_missing": ["skill3", "skill4", ...],
        "partially_covered": [{"skill": "skillname", "gap": "what's missing"}]
    },
    "matching_skills": ["skill1", "skill2", ...],
    "recommendations": {
        "immediate_actions": ["action1", "action2", ...],
        "learning_path": [{"skill": "skillname", "priority": "high/medium/low", "timeframe": "weeks", "resources": ["resource1", "resource2"]}],
        "certification_suggestions": ["cert1", "cert2", ...]
    },
    "fit_score": 85,
    "key_strengths": ["strength1", "strength2", ...],
    "improvement_areas": ["area1", "area2", ...]
}"""
                        },
                        {
                            "role": "user",
                            "content": f"""Analyze this resume against the job requirements{company_context}:

RESUME:
{resume_text}

JOB REQUIREMENTS:
{job_description}

Provide a comprehensive skill gap analysis with specific, actionable recommendations."""
                        }
                    ]
                    
                    response = make_chat_completion(messages, model=selected_model, max_tokens=2000)
                    analysis_text = response["choices"][0]["message"]["content"]
                    
                    # Try to parse JSON response
                    try:
                        # Extract JSON from the response
                        import re
                        json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                        if json_match:
                            analysis = json.loads(json_match.group())
                        else:
                            # Fallback if JSON parsing fails
                            st.error("Could not parse structured analysis. Showing raw response:")
                            st.text(analysis_text)
                            return
                    except:
                        # If JSON parsing fails, show raw analysis
                        st.markdown("### 📊 Skill Gap Analysis Results")
                        st.markdown(analysis_text)
                        return
                    
                    # Display structured results
                    st.markdown("### 📊 Skill Gap Analysis Results")
                    
                    # Overall fit score
                    _, col2, _ = st.columns(3)
                    with col2:
                        fit_score = analysis.get("fit_score", 0)
                        fig = go.Figure(go.Indicator(
                            mode="gauge+number+delta",
                            value=fit_score,
                            title={'text': "Overall Fit Score"},
                            gauge={
                                'axis': {'range': [None, 100]},
                                'bar': {'color': "darkblue"},
                                'steps': [
                                    {'range': [0, 60], 'color': "lightgray"},
                                    {'range': [60, 80], 'color': "yellow"},
                                    {'range': [80, 100], 'color': "green"}
                                ],
                                'threshold': {
                                    'line': {'color': "red", 'width': 4},
                                    'thickness': 0.75,
                                    'value': 90
                                }
                            }
                        ))
                        fig.update_layout(height=300)
                        st.plotly_chart(fig, use_container_width=True)
                    
                    # Skill gaps section
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### ❌ Missing Skills")
                        skill_gaps = analysis.get("skill_gaps", {})
                        
                        if skill_gaps.get("critical_missing"):
                            st.markdown("**🔴 Critical Missing:**")
                            for skill in skill_gaps["critical_missing"]:
                                st.markdown(f"• {skill}")
                        
                        if skill_gaps.get("nice_to_have_missing"):
                            st.markdown("**🟡 Nice to Have:**")
                            for skill in skill_gaps["nice_to_have_missing"]:
                                st.markdown(f"• {skill}")
                        
                        if skill_gaps.get("partially_covered"):
                            st.markdown("**🟠 Partially Covered:**")
                            for item in skill_gaps["partially_covered"]:
                                st.markdown(f"• **{item['skill']}**: {item['gap']}")
                    
                    with col2:
                        st.markdown("#### ✅ Matching Skills")
                        matching_skills = analysis.get("matching_skills", [])
                        if matching_skills:
                            for skill in matching_skills:
                                st.markdown(f"• {skill}")
                        else:
                            st.markdown("*No matching skills identified*")
                    
                    # Key strengths and improvement areas
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### 💪 Key Strengths")
                        strengths = analysis.get("key_strengths", [])
                        for strength in strengths:
                            st.markdown(f"• {strength}")
                    
                    with col2:
                        st.markdown("#### 📈 Improvement Areas")
                        improvements = analysis.get("improvement_areas", [])
                        for area in improvements:
                            st.markdown(f"• {area}")
                    
                    # Recommendations
                    st.markdown("#### 🎯 Recommendations")
                    recommendations = analysis.get("recommendations", {})
                    
                    if recommendations.get("immediate_actions"):
                        st.markdown("**Immediate Actions:**")
                        for action in recommendations["immediate_actions"]:
                            st.markdown(f"• {action}")
                    
                    if recommendations.get("learning_path"):
                        st.markdown("**📚 Learning Path:**")
                        learning_df = []
                        for item in recommendations["learning_path"]:
                            priority_emoji = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(item.get("priority", "medium"), "🟡")
                            learning_df.append({
                                "Skill": item.get("skill", ""),
                                "Priority": f"{priority_emoji} {item.get('priority', 'medium').title()}",
                                "Timeframe": item.get("timeframe", ""),
                                "Resources": ", ".join(item.get("resources", []))
                            })
                        
                        if learning_df:
                            st.dataframe(learning_df, use_container_width=True)
                    
                    if recommendations.get("certification_suggestions"):
                        st.markdown("**🏆 Suggested Certifications:**")
                        for cert in recommendations["certification_suggestions"]:
                            st.markdown(f"• {cert}")
                
                except Exception as e:
                    st.error(f"Error during analysis: {e}")


def company_culture_page():
    """Company Culture & Skill Requirements page"""
    
    st.title("🏢 Company Culture & Skills Analysis")
    st.markdown("**Research company culture and get detailed skill requirements for specific roles**")
    
    # Model Selection for this page
    with st.sidebar:
        st.subheader("⚙️ Analysis Settings")
        selected_model = st.selectbox(
            "Select AI Model:",
            options=list(AVAILABLE_MODELS.keys()),
            format_func=lambda x: f"{AVAILABLE_MODELS[x]['name']} - {AVAILABLE_MODELS[x]['description']}",
            key="culture_model"
        )
    
    # Input section
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🎯 Company Information")
        company_name = st.text_input("Company Name*:", placeholder="e.g., Google, Microsoft, Amazon")
        job_title = st.text_input("Job Title*:", placeholder="e.g., Software Engineer, Data Scientist")
        location = st.text_input("Location (optional):", placeholder="e.g., San Francisco, Remote")
        
    with col2:
        st.subheader("📋 Additional Context")
        industry = st.selectbox(
            "Industry:",
            ["Technology", "Finance", "Healthcare", "Education", "Retail", "Manufacturing",
             "Consulting", "Media", "Government", "Non-profit", "Other"],
            key="industry_select"
        )
        company_size = st.selectbox(
            "Company Size:",
            ["Startup (1-50)", "Small (51-200)", "Medium (201-1000)", "Large (1001-5000)", "Enterprise (5000+)", "Unknown"],
            key="size_select"
        )
        experience_level = st.selectbox(
            "Experience Level:",
            ["Entry Level", "Junior (1-3 years)", "Mid-level (3-7 years)", "Senior (7-12 years)", "Lead/Principal (12+ years)"],
            key="exp_level"
        )
    
    # Optional job description
    st.subheader("📄 Job Description (Optional)")
    job_description = st.text_area(
        "Paste job description for more detailed analysis:",
        height=150,
        placeholder="Optional: Paste the job description here for more specific analysis..."
    )
    
    if st.button("🔍 Analyze Company & Requirements", type="primary"):
        if not company_name.strip() or not job_title.strip():
            st.error("Please provide at least the company name and job title!")
        else:
            with st.spinner("Analyzing company culture and skill requirements..."):
                try:
                    # Prepare the analysis prompt
                    location_context = f" in {location}" if location else ""
                    job_desc_context = f"\n\nJOB DESCRIPTION:\n{job_description}" if job_description.strip() else ""
                    
                    messages = [
                        {
                            "role": "system",
                            "content": """You are an expert company researcher and career analyst. Provide comprehensive insights about company culture, values, and detailed skill requirements for specific roles.

Provide your analysis in the following JSON format:
{
    "company_overview": {
        "mission": "company mission/vision",
        "values": ["value1", "value2", "value3"],
        "culture_highlights": ["culture1", "culture2", "culture3"],
        "work_environment": "description of work environment",
        "growth_stage": "startup/growth/mature/enterprise",
        "reputation": "industry reputation and notable achievements"
    },
    "role_requirements": {
        "technical_skills": [
            {"skill": "skillname", "importance": "critical/important/nice-to-have", "proficiency_level": "beginner/intermediate/advanced/expert"}
        ],
        "soft_skills": ["skill1", "skill2", "skill3"],
        "experience_requirements": "detailed experience expectations",
        "education_requirements": ["requirement1", "requirement2"],
        "certifications": ["cert1", "cert2"],
        "salary_range": "estimated salary range",
        "career_progression": "typical career path"
    },
    "interview_insights": {
        "interview_style": "description of interview process",
        "common_questions": ["question1", "question2", "question3"],
        "technical_assessment": "description of technical evaluation",
        "culture_fit_factors": ["factor1", "factor2", "factor3"]
    },
    "preparation_tips": {
        "research_focus": ["area1", "area2", "area3"],
        "skill_priorities": ["priority1", "priority2", "priority3"],
        "networking_opportunities": ["opportunity1", "opportunity2"],
        "application_tips": ["tip1", "tip2", "tip3"]
    }
}"""
                        },
                        {
                            "role": "user",
                            "content": f"""Analyze the company culture and provide detailed skill requirements for:

Company: {company_name}
Position: {job_title}{location_context}
Industry: {industry}
Company Size: {company_size}
Experience Level: {experience_level}{job_desc_context}

Provide comprehensive insights about company culture, detailed skill requirements, interview process, and preparation tips."""
                        }
                    ]
                    
                    response = make_chat_completion(messages, model=selected_model, max_tokens=2500)
                    analysis_text = response["choices"][0]["message"]["content"]
                    
                    # Try to parse JSON response
                    try:
                        # Extract JSON from the response
                        import re
                        json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                        if json_match:
                            analysis = json.loads(json_match.group())
                        else:
                            # Fallback if JSON parsing fails
                            st.error("Could not parse structured analysis. Showing raw response:")
                            st.markdown(analysis_text)
                            return
                    except:
                        # If JSON parsing fails, show raw analysis
                        st.markdown("### 🏢 Company Culture & Requirements Analysis")
                        st.markdown(analysis_text)
                        return
                    
                    # Display structured results
                    st.markdown(f"### 🏢 {company_name} - {job_title} Analysis")
                    
                    # Company Overview
                    company_overview = analysis.get("company_overview", {})
                    if company_overview:
                        st.markdown("#### 🌟 Company Overview")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            if company_overview.get("mission"):
                                st.markdown(f"**Mission:** {company_overview['mission']}")
                            
                            if company_overview.get("values"):
                                st.markdown("**Core Values:**")
                                for value in company_overview["values"]:
                                    st.markdown(f"• {value}")
                        
                        with col2:
                            if company_overview.get("work_environment"):
                                st.markdown(f"**Work Environment:** {company_overview['work_environment']}")
                            
                            if company_overview.get("growth_stage"):
                                st.markdown(f"**Growth Stage:** {company_overview['growth_stage']}")
                        
                        if company_overview.get("culture_highlights"):
                            st.markdown("**Culture Highlights:**")
                            for highlight in company_overview["culture_highlights"]:
                                st.markdown(f"• {highlight}")
                        
                        if company_overview.get("reputation"):
                            st.info(f"**Industry Reputation:** {company_overview['reputation']}")
                    
                    # Role Requirements
                    role_requirements = analysis.get("role_requirements", {})
                    if role_requirements:
                        st.markdown("#### 💼 Role Requirements")
                        
                        # Technical Skills
                        if role_requirements.get("technical_skills"):
                            st.markdown("**🔧 Technical Skills:**")
                            
                            # Create DataFrame for technical skills
                            tech_skills_df = []
                            for skill in role_requirements["technical_skills"]:
                                importance_emoji = {
                                    "critical": "🔴",
                                    "important": "🟡",
                                    "nice-to-have": "🟢"
                                }.get(skill.get("importance", "important").lower(), "🟡")
                                
                                tech_skills_df.append({
                                    "Skill": skill.get("skill", ""),
                                    "Importance": f"{importance_emoji} {skill.get('importance', 'important').title()}",
                                    "Level Required": skill.get("proficiency_level", "intermediate").title()
                                })
                            
                            if tech_skills_df:
                                st.dataframe(tech_skills_df, use_container_width=True)
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Soft Skills
                            if role_requirements.get("soft_skills"):
                                st.markdown("**🤝 Soft Skills:**")
                                for skill in role_requirements["soft_skills"]:
                                    st.markdown(f"• {skill}")
                            
                            # Education Requirements
                            if role_requirements.get("education_requirements"):
                                st.markdown("**🎓 Education Requirements:**")
                                for req in role_requirements["education_requirements"]:
                                    st.markdown(f"• {req}")
                        
                        with col2:
                            # Experience Requirements
                            if role_requirements.get("experience_requirements"):
                                st.markdown(f"**💼 Experience:** {role_requirements['experience_requirements']}")
                            
                            # Certifications
                            if role_requirements.get("certifications"):
                                st.markdown("**🏆 Recommended Certifications:**")
                                for cert in role_requirements["certifications"]:
                                    st.markdown(f"• {cert}")
                        
                        # Salary and Career Progression
                        col1, col2 = st.columns(2)
                        with col1:
                            if role_requirements.get("salary_range"):
                                st.success(f"💰 **Salary Range:** {role_requirements['salary_range']}")
                        
                        with col2:
                            if role_requirements.get("career_progression"):
                                st.info(f"📈 **Career Path:** {role_requirements['career_progression']}")
                    
                    # Interview Insights
                    interview_insights = analysis.get("interview_insights", {})
                    if interview_insights:
                        st.markdown("#### 🎤 Interview Insights")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if interview_insights.get("interview_style"):
                                st.markdown(f"**Interview Style:** {interview_insights['interview_style']}")
                            
                            if interview_insights.get("technical_assessment"):
                                st.markdown(f"**Technical Assessment:** {interview_insights['technical_assessment']}")
                        
                        with col2:
                            if interview_insights.get("culture_fit_factors"):
                                st.markdown("**Culture Fit Factors:**")
                                for factor in interview_insights["culture_fit_factors"]:
                                    st.markdown(f"• {factor}")
                        
                        if interview_insights.get("common_questions"):
                            st.markdown("**Common Interview Questions:**")
                            for i, question in enumerate(interview_insights["common_questions"], 1):
                                st.markdown(f"{i}. {question}")
                    
                    # Preparation Tips
                    preparation_tips = analysis.get("preparation_tips", {})
                    if preparation_tips:
                        st.markdown("#### 🎯 Preparation Tips")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if preparation_tips.get("research_focus"):
                                st.markdown("**🔍 Research Focus:**")
                                for focus in preparation_tips["research_focus"]:
                                    st.markdown(f"• {focus}")
                            
                            if preparation_tips.get("skill_priorities"):
                                st.markdown("**📚 Skill Priorities:**")
                                for priority in preparation_tips["skill_priorities"]:
                                    st.markdown(f"• {priority}")
                        
                        with col2:
                            if preparation_tips.get("networking_opportunities"):
                                st.markdown("**🤝 Networking Opportunities:**")
                                for opportunity in preparation_tips["networking_opportunities"]:
                                    st.markdown(f"• {opportunity}")
                            
                            if preparation_tips.get("application_tips"):
                                st.markdown("**✅ Application Tips:**")
                                for tip in preparation_tips["application_tips"]:
                                    st.markdown(f"• {tip}")
                
                except Exception as e:
                    st.error(f"Error during analysis: {e}")


def resume_comparison_page():
    """Resume comparison page for comparing two PDFs"""
    st.title("⚖️ Resume Comparison")
    st.markdown("Upload two resumes to compare them side-by-side. Optionally add a job description for targeted comparison.")

    # Model Selection for this page
    with st.sidebar:
        st.markdown("### 🤖 Model Selection")
        model_choice = st.selectbox(
            "Choose Analysis Model",
            options=list(AVAILABLE_MODELS.keys()),
            format_func=lambda x: AVAILABLE_MODELS[x]["name"],
            index=list(AVAILABLE_MODELS.keys()).index(DEFAULT_MODEL),
            help="Select the AI model to use for comparison analysis.",
            key="comparison_model"
        )
        st.caption(f"💡 {AVAILABLE_MODELS[model_choice]['description']}")

    # File uploads
    st.subheader("📋 Upload Resumes to Compare")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("##### Resume A")
        resume_a_file = st.file_uploader(
            "Upload first resume (PDF/DOCX/TXT)",
            type=["pdf", "docx", "txt"],
            key="resume_a"
        )
    
    with col2:
        st.markdown("##### Resume B") 
        resume_b_file = st.file_uploader(
            "Upload second resume (PDF/DOCX/TXT)",
            type=["pdf", "docx", "txt"],
            key="resume_b"
        )

    # Optional job description
    st.subheader("🎯 Job Description (Optional)")
    st.markdown("*Add a job description to get targeted comparison based on specific requirements.*")
    jd_text = st.text_area(
        "Paste job description here (optional but recommended)",
        height=150,
        key="comparison_jd"
    )

    # Comparison button
    compare_btn = st.button("⚖️ Compare Resumes", type="primary")

    if compare_btn:
        if not resume_a_file or not resume_b_file:
            st.error("Please upload both resumes before comparing.")
        else:
            # Extract text from both files
            with st.spinner("Extracting text from resumes..."):
                try:
                    resume_a_text = extract_text_from_file(resume_a_file)
                    resume_b_text = extract_text_from_file(resume_b_file)
                except Exception as e:
                    st.error(f"Failed to extract text from files: {e}")
                    return

            if resume_a_text and resume_b_text:
                st.info("Analyzing and comparing resumes with AI...")
                try:
                    comparison_data = call_for_comparison(
                        resume_a_text,
                        resume_b_text,
                        jd_text or "",
                        model_choice
                    )
                except Exception as e:
                    st.error(f"Error from API: {e}")
                    return

                if comparison_data:
                    # Display comparison results
                    st.subheader("📊 Comparison Results")
                    
                    # Winner announcement
                    winner = comparison_data.get("winner", "Tie")
                    if winner != "Tie":
                        st.success(f"🏆 **Winner: Resume {winner}**")
                    else:
                        st.info("🤝 **Result: It's a tie!**")
                    
                    # Overall comparison
                    st.markdown("### 📝 Overall Analysis")
                    st.write(comparison_data.get("overall_comparison", "No comparison available"))
                    
                    # Scores visualization
                    if comparison_data.get("resume_a_score") and comparison_data.get("resume_b_score"):
                        comparison_chart = create_comparison_visualization(comparison_data)
                        st.plotly_chart(comparison_chart, use_container_width=True)
                    
                    # Side-by-side strengths
                    st.markdown("### 💪 Strengths Comparison")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### Resume A Strengths")
                        for strength in comparison_data.get("resume_a_strengths", []):
                            st.success(f"• {strength}")
                        
                        score_a = comparison_data.get("resume_a_score", "N/A")
                        st.metric("Resume A Score", f"{score_a}/100" if score_a != "N/A" else score_a)
                    
                    with col2:
                        st.markdown("#### Resume B Strengths") 
                        for strength in comparison_data.get("resume_b_strengths", []):
                            st.success(f"• {strength}")
                            
                        score_b = comparison_data.get("resume_b_score", "N/A")
                        st.metric("Resume B Score", f"{score_b}/100" if score_b != "N/A" else score_b)
                    
                    # Recommendations
                    st.markdown("### 🎯 Improvement Recommendations")
                    for i, recommendation in enumerate(comparison_data.get("recommendations", []), 1):
                        st.write(f"**{i}.** {recommendation}")
                    
                    # Resume previews
                    st.markdown("### 📄 Resume Previews")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### Resume A Content")
                        with st.expander("View Resume A", expanded=False):
                            st.text_area(
                                "Resume A Text",
                                resume_a_text[:2000] + "..." if len(resume_a_text) > 2000 else resume_a_text,
                                height=300,
                                key="preview_a"
                            )
                    
                    with col2:
                        st.markdown("#### Resume B Content")
                        with st.expander("View Resume B", expanded=False):
                            st.text_area(
                                "Resume B Text",
                                resume_b_text[:2000] + "..." if len(resume_b_text) > 2000 else resume_b_text,
                                height=300,
                                key="preview_b"
                            )
    
    # Tips section
    st.markdown("---")
    st.markdown("""
    **💡 Comparison Tips:**
    - Upload resumes in similar formats for better comparison
    - Include a job description for more targeted analysis  
    - The AI considers ATS compatibility, content quality, and job alignment
    - Use this to identify which resume performs better for specific roles
    """)


# --------------------------
# Main App
# --------------------------
def main():
    st.set_page_config(page_title="AI Resume Analyzer", layout="wide")

    # Sidebar for navigation
    st.sidebar.title("🔍 AI Resume Analyzer")
    page = st.sidebar.selectbox(
        "Navigate to:",
        ["📊 Resume Analysis", "🎯 Skill Gap Analysis", "🏢 Company Culture & Skills", "⚖️ Resume Comparison"],
        index=0
    )

    # Page routing
    if page == "📊 Resume Analysis":
        resume_analysis_page()
    elif page == "🎯 Skill Gap Analysis":
        skill_gap_analysis_page()
    elif page == "🏢 Company Culture & Skills":
        company_culture_page()
    elif page == "⚖️ Resume Comparison":
        resume_comparison_page()


if __name__ == "__main__":
    main()
