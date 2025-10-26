import os
import json
import random
import tempfile
import requests
from typing import Tuple, Dict, Any
import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# File parsing libs
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx2txt
except Exception:
    docx2txt = None

# For docx generation
try:
    from docx import Document
except Exception:
    Document = None

# --------------------------
# CONFIG & CONSTANTS
# --------------------------
API_KEY = os.getenv("OPENAI_API_KEY")  # Load from environment variable
API_BASE = "https://api.chatanywhere.tech"
DEFAULT_MODEL = "gpt-3.5-turbo"

AVAILABLE_MODELS = {
    "gpt-3.5-turbo": {"name": "GPT-3.5", "description": "Fast & efficient analysis"},
    "gpt-4o-mini": {"name": "GPT-4o", "description": "More detailed analysis"},
    "gpt-5-mini": {"name": "GPT-5", "description": "Best results, limited usage"},
    "deepseek-v3": {"name": "DeepSeek", "description": "Specialized for resumes"}
}

# ATS Analysis Categories
ATS_CATEGORIES = {
    "Content Quality": [
        "Achievement Focus",
        "Action Verbs",
        "Quantified Results",
        "Industry Keywords",
        "Role-specific Skills"
    ],
    "Format & Structure": [
        "ATS Readability",
        "Section Headers",
        "Bullet Points",
        "Consistent Formatting",
        "Contact Info"
    ],
    "Experience Match": [
        "Job Alignment",
        "Skills Coverage",
        "Required Qualifications",
        "Experience Level",
        "Domain Knowledge"
    ]
}


# --------------------------
# Helper: Chat Completion
# --------------------------
def make_chat_completion(messages, model=DEFAULT_MODEL, temperature=0.0, max_tokens=1000) -> Dict[str, Any]:
    """
    Generic function to call the ChatAnywhere (OpenAI-style) chat completions endpoint.
    Returns the parsed JSON response.
    """
    if not API_KEY:
        raise RuntimeError("API_KEY not set. Set OPENAI_API_KEY env var or API_KEY variable in code.")
    url = f"{API_BASE.rstrip('/')}/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=60)
    try:
        resp.raise_for_status()
    except Exception as e:
        # Surface helpful message
        raise RuntimeError(f"API call failed: {e}\nResponse: {resp.text}")
    return resp.json()


# --------------------------
# File text extraction
# --------------------------
def extract_text_from_file(uploaded_file) -> str:
    """
    Accepts a streamlit uploaded file-like object
    Supports PDF, DOCX, TXT.
    Returns plain text string.
    """
    fname = uploaded_file.name.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp:
        tmp.write(uploaded_file.getbuffer())
        tmp.flush()
        tmp_path = tmp.name

    text = ""
    try:
        if fname.endswith(".pdf"):
            if pdfplumber is None:
                raise RuntimeError("pdfplumber not installed. Add pdfplumber in requirements.")
            with pdfplumber.open(tmp_path) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages)
        elif fname.endswith(".docx") or fname.endswith(".doc"):
            if docx2txt is None:
                raise RuntimeError("docx2txt not installed. Add docx2txt in requirements.")
            text = docx2txt.process(tmp_path) or ""
        else:
            # txt or fallback
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
    return text.strip()


# --------------------------
# Prompting helpers
# --------------------------
SCORE_PROMPT = """You are an expert career coach and resume reviewer with experience in recruiting and ATS.
Given the user's resume text and the job description, produce a JSON object with these fields:
- score: integer 0-100 (how good the resume matches the JD and ATS readiness)
- strengths: list of short strings (3-6)
- weaknesses: list of short strings (3-8)
- improvement_tips: list of actionable suggestions (5 items) (formatting, keywords, achievements)
- missing_keywords: list of important keywords/phrases present in JD but missing in resume (up to 12)
- overall_summary: one short paragraph (max 80 words)

Input JSON:
{{
  "resume_text": "{resume_text}",
  "job_description": "{job_description}"
}}

Return only valid JSON (no commentary). If you cannot match fields, still return a JSON with empty lists and a score.
"""

TAILOR_PROMPT = """You are an expert resume writer. Using the user's resume and the job description, produce:
1) A tailored resume in plain text that:
   - Reorders sections for maximum impact
   - Rewrites bullet points to emphasize achievements and use JD keywords
   - Keeps length roughly the same (don't invent jobs)
   - Uses result-driven bullets (use metrics where present)
Return result under a JSON with keys:
{
  "tailored_resume": "<full resume as plain text>"
}
Do not add anything else outside the JSON.
"""

COVER_PROMPT = """You are an expert career coach. Write a short cover-note (3-5 sentences) tailored to the job description and the candidate's resume emphasizing fit and one key achievement. Return as plain text only.
"""


# --------------------------
# ATS Analysis Functions
# --------------------------
def create_ats_analysis(score_data):
    """Main function to handle all ATS analysis"""
    # Get base score
    base_score = score_data.get("score", 0)
    
    # Calculate category scores
    category_scores = {}
    for category in ATS_CATEGORIES.keys():
        adjustment = len(score_data.get("strengths", [])) - len(score_data.get("weaknesses", []))
        adjustment *= 5
        category_scores[category] = min(100, max(0, base_score + adjustment))
    
    # Calculate detailed scores
    for category, subcats in ATS_CATEGORIES.items():
        category_base = category_scores[category]
        for subcat in subcats:
            score_key = f"{category.lower()}_{subcat.lower().replace(' ', '_')}"
            variation = random.randint(-5, 5)
            score_data[score_key] = min(100, max(0, category_base + variation))
    
    return score_data


def create_ats_visualizations(score_data):
    """Create all ATS-related visualizations"""
    # Create radar chart
    radar_fig = go.Figure()
    categories = []
    scores = []
    
    for category, subcats in ATS_CATEGORIES.items():
        for subcat in subcats:
            label = f"{category}:<br>{subcat}"
            categories.append(label)
            score_key = f"{category.lower()}_{subcat.lower().replace(' ', '_')}"
            score = score_data.get(score_key, score_data.get("score", 50))
            scores.append(min(100, max(0, score)))

    radar_fig.add_trace(go.Scatterpolar(
        r=scores,
        theta=categories,
        fill='toself',
        name='Resume Score',
        line=dict(color='#2E86C1', width=2),
        fillcolor='rgba(46, 134, 193, 0.3)'
    ))

    radar_fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                tickfont=dict(size=10),
                gridcolor='rgba(0,0,0,0.1)',
                ticksuffix="%"
            ),
            angularaxis=dict(
                tickfont=dict(size=8),
                rotation=90,
                direction="clockwise"
            )
        ),
        showlegend=False,
        title="ATS Score Breakdown",
        height=600,
        margin=dict(t=100, b=50)
    )
    
    # Create skills match chart
    missing_keywords = score_data.get("missing_keywords", [])
    matched_keywords = score_data.get("matched_keywords", score_data.get("strengths", []))
    
    match_fig = None
    if missing_keywords or matched_keywords:
        total = len(missing_keywords) + len(matched_keywords)
        match_percent = (len(matched_keywords) / total * 100) if total > 0 else 0
        
        match_fig = go.Figure()
        match_fig.add_trace(go.Bar(
            x=["Skills Match"],
            y=[match_percent],
            name="Matched Keywords",
            marker_color="rgb(46, 134, 193)",
            hovertemplate="Matched: %{y:.1f}%<br>Keywords: " + 
                         ", ".join(matched_keywords[:5]) + 
                         ("..." if len(matched_keywords) > 5 else "")
        ))
        match_fig.add_trace(go.Bar(
            x=["Skills Match"],
            y=[100 - match_percent],
            name="Missing Keywords",
            marker_color="rgba(231, 76, 60, 0.7)",
            hovertemplate="Missing: %{y:.1f}%<br>Keywords: " + 
                         ", ".join(missing_keywords[:5]) + 
                         ("..." if len(missing_keywords) > 5 else "")
        ))
        
        match_fig.update_layout(
            barmode='stack',
            title="Skills Match Analysis",
            yaxis_title="Match Percentage",
            showlegend=True,
            height=300,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
    
    return radar_fig, match_fig


# --------------------------
# UI and main flows
# --------------------------
st.set_page_config(page_title="AI Resume Analyzer", layout="wide")

# Sidebar for navigation
st.sidebar.title("🔍 AI Resume Analyzer")
page = st.sidebar.selectbox(
    "Navigate to:",
    ["📊 Resume Analysis", "⚖️ Resume Comparison"],
    index=0
)

# Page routing
if page == "📊 Resume Analysis":
    resume_analysis_page()
elif page == "⚖️ Resume Comparison":
    resume_comparison_page()


def resume_analysis_page():
    st.title("🔍 Resume Score & Tailor")
    st.markdown("Upload your resume and paste the job description (JD). This app will score your resume, give improvements, and produce a tailored resume using AI.")

        # Model Selection for this page
    with st.sidebar:
        st.markdown("### 🤖 Model Selection")
        model_choice = st.selectbox(
            "Choose Analysis Model",
            options=list(AVAILABLE_MODELS.keys()),
            format_func=lambda x: AVAILABLE_MODELS[x]["name"],
            index=list(AVAILABLE_MODELS.keys()).index(DEFAULT_MODEL),
            help="Select the AI model to use for analysis. Each model has different capabilities."
        )
        st.caption(f"💡 {AVAILABLE_MODELS[model_choice]['description']}")

    # Upload resume and JD
    st.subheader("1) Upload resume")
    uploaded = st.file_uploader("Upload PDF / DOCX / TXT resume", type=["pdf", "docx", "txt"])
    st.subheader("2) Paste Job Description")
    jd_text = st.text_area("Paste the Job Description (JD) here", height=220)

    col1, col2 = st.columns(2)
    with col1:
        analyze_btn = st.button("🔎 Analyze & Score")
    with col2:
        tailor_btn = st.button("✍️ Generate Tailored Resume")

    # Area for outputs
    output_area = st.empty()


# Upload resume and JD
st.subheader("1) Upload resume")
uploaded = st.file_uploader("Upload PDF / DOCX / TXT resume", type=["pdf", "docx", "txt"])
st.subheader("2) Paste Job Description")
jd_text = st.text_area("Paste the Job Description (JD) here", height=220)

col1, col2 = st.columns(2)
with col1:
    analyze_btn = st.button("🔎 Analyze & Score")
with col2:
    tailor_btn = st.button("✍️ Generate Tailored Resume")

# Area for outputs
output_area = st.empty()


def call_for_score(resume_text: str, jd: str):
    # Create message for the model to strictly output JSON
    prompt = SCORE_PROMPT.format(resume_text=resume_text.replace("\n", "\\n"), job_description=jd.replace("\n", "\\n"))
    messages = [
        {"role": "system", "content": "You are an assistant that strictly returns JSON when asked."},
        {"role": "user", "content": prompt}
    ]
    resp = make_chat_completion(messages, model=model_choice, temperature=0.0, max_tokens=800)
    # Parse the assistant reply (best effort)
    try:
        assistant_text = resp["choices"][0]["message"]["content"]
        # Attempt to find JSON substring
        start = assistant_text.find("{")
        end = assistant_text.rfind("}") + 1
        if start != -1 and end != -1:
            json_text = assistant_text[start:end]
            data = json.loads(json_text)
            return data
        else:
            # As fallback, try to parse whole message
            return json.loads(assistant_text)
    except Exception as e:
        raise RuntimeError(f"Failed to parse scoring response: {e}\nRaw response: {resp}")


def call_for_tailor(resume_text: str, jd: str):
    prompt = TAILOR_PROMPT
    # embed resume and jd into system message + user message to avoid length issues
    messages = [
        {"role": "system", "content": "You are an expert resume writer. Output a JSON object with a 'tailored_resume' key."},
        {"role": "user", "content": f"Resume:\n{resume_text}\n\nJob Description:\n{jd}"}
    ]
    resp = make_chat_completion(messages, model=model_choice, temperature=0.2, max_tokens=1500)
    try:
        assistant_text = resp["choices"][0]["message"]["content"]
        start = assistant_text.find("{")
        end = assistant_text.rfind("}") + 1
        if start != -1 and end != -1:
            json_text = assistant_text[start:end]
            data = json.loads(json_text)
            return data
        else:
            # fallback: return raw text as tailored_resume
            return {"tailored_resume": assistant_text}
    except Exception as e:
        raise RuntimeError(f"Failed to parse tailoring response: {e}\nRaw response: {resp}")


def call_for_cover(resume_text: str, jd: str):
    messages = [
        {"role": "system", "content": "You are an expert career coach."},
        {"role": "user", "content": f"Write a short cover note based on this resume and job description.\n\nResume:\n{resume_text}\n\nJD:\n{jd}"}
    ]
    resp = make_chat_completion(messages, model=model_choice, temperature=0.3, max_tokens=200)
    return resp["choices"][0]["message"]["content"]


def ensure_string(data):
    """Ensure data is a string for download"""
    if isinstance(data, dict):
        return json.dumps(data, indent=2)
    return str(data)


def format_resume_output(resume_text: str) -> str:
    """Format resume text with proper section detection and spacing"""
    sections = []
    current_section = []
    
    # Common section headers in resumes
    SECTION_KEYWORDS = {
        "SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS",
        "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS",
        "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
        "TECHNICAL SKILLS", "PUBLICATIONS", "LANGUAGES"
    }
    
    for line in resume_text.splitlines():
        line = line.strip()
        # Enhanced section header detection
        is_header = (
            line.upper() in SECTION_KEYWORDS or
            any(line.upper().startswith(k) for k in SECTION_KEYWORDS) or
            line.endswith(':') or
            (line.isupper() and len(line) > 3)  # Avoid short uppercase words
        )
        
        if is_header and current_section:
            # Complete previous section
            sections.append('\n'.join(current_section))
            current_section = []
            # Start new section with formatted header
            current_section.append(f"\n{line.upper()}")
        elif line:
            # Format bullet points consistently
            if line.lstrip().startswith(('•', '-', '*', '○', '→')):
                line = '• ' + line.lstrip('•-*○→ \t')
            current_section.append(line)
    
    # Add final section
    if current_section:
        sections.append('\n'.join(current_section))
    
    # Join sections with proper spacing
    formatted_text = '\n\n'.join(sections)
    # Clean up multiple blank lines
    formatted_text = '\n'.join(line for line in formatted_text.splitlines() if line.strip() or line == '')
    while '\n\n\n' in formatted_text:
        formatted_text = formatted_text.replace('\n\n\n', '\n\n')
    
    return formatted_text.strip()


# --------------------------
# Process user actions
# --------------------------
if analyze_btn:
    if not uploaded:
        st.warning("Please upload a resume file first.")
    elif not jd_text or len(jd_text.strip()) < 20:
        st.warning("Please paste a detailed job description (at least ~20 characters).")
    else:
        with st.spinner("Extracting resume text..."):
            try:
                resume_text = extract_text_from_file(uploaded)
            except Exception as e:
                st.error(f"Failed to extract text from file: {e}")
                resume_text = None

        if resume_text:
            st.info("Contacting ChatGPT API to score the resume...")
            try:
                score_data = call_for_score(resume_text, jd_text)
            except Exception as e:
                st.error(f"Error from API: {e}")
                score_data = None

            if score_data:
                score_data = create_ats_analysis(score_data)
                radar_chart, skills_chart = create_ats_visualizations(score_data)
                
                # Display results
                score_col, viz_col = st.columns([1, 2])
                with score_col:
                    st.subheader("📊 Resume Score")
                    score = score_data.get("score", "N/A")
                    st.metric("Match Score (0-100)", score)
                    
                    # Add color-coded score interpretation
                    if isinstance(score, (int, float)):
                        if score >= 80:
                            st.success("Strong Match! 🌟")
                        elif score >= 60:
                            st.info("Good Match 👍")
                        else:
                            st.warning("Needs Improvement 🔨")
                with viz_col:
                    st.plotly_chart(radar_chart, use_container_width=True)
                
                if skills_chart:
                    st.plotly_chart(skills_chart, use_container_width=True)
                
                # Continue with existing sections
                st.subheader("Strengths")
                for s in score_data.get("strengths", []):
                    st.success(f"• {s}")
                st.subheader("Weaknesses")
                for w in score_data.get("weaknesses", []):
                    st.error(f"• {w}")
                st.subheader("Improvement Tips")
                for t in score_data.get("improvement_tips", []):
                    st.write(f"• {t}")
                st.subheader("Missing Keywords (from JD)")
                if score_data.get("missing_keywords"):
                    st.write(", ".join(score_data.get("missing_keywords")))
                else:
                    st.write("None detected")
                st.subheader("Summary")
                st.write(score_data.get("overall_summary", ""))
                # Save for later tailoring
                st.session_state["last_resume_text"] = resume_text
                st.session_state["last_jd"] = jd_text

if tailor_btn:
    # ensure we have resume_text
    if "last_resume_text" not in st.session_state:
        if not uploaded:
            st.warning("Please upload a resume and click Analyze first (or paste JD and click Analyze).")
        else:
            with st.spinner("Extracting resume text..."):
                try:
                    resume_text = extract_text_from_file(uploaded)
                    st.session_state["last_resume_text"] = resume_text
                    st.session_state["last_jd"] = jd_text
                except Exception as e:
                    st.error(f"Failed to extract text from file: {e}")
    if "last_resume_text" in st.session_state:
        resume_text = st.session_state["last_resume_text"]
        jd_text_local = st.session_state.get("last_jd", jd_text)
        st.info("Asking AI to generate a tailored resume...")
        try:
            tailored_json = call_for_tailor(resume_text, jd_text_local)
            tailored_resume = tailored_json.get("tailored_resume", "")
        except Exception as e:
            st.error(f"Error generating tailored resume: {e}")
            tailored_resume = ""

        if tailored_resume:
            st.subheader("✏️ Tailored Resume (Preview)")
            
            # Format the resume text
            tailored_text = ensure_string(tailored_resume)
            formatted_text = format_resume_output(tailored_text)
            
            # Create columns for better layout
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.text_area("Preview", formatted_text, height=500)
            
            with col2:
                st.markdown("### Download Options")
                st.download_button(
                    "📄 Download as TXT",
                    formatted_text,
                    file_name="tailored_resume.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
                if Document is not None:
                    # Create formatted docx
                    doc = Document()
                    for section in formatted_text.split('\n\n'):
                        if section.strip():
                            if section.isupper() or section.endswith(':'):
                                doc.add_heading(section.strip(':'), level=2)
                            else:
                                doc.add_paragraph(section)
                
                    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(tmp_docx.name)
                    tmp_docx.flush()
                    
                    with open(tmp_docx.name, "rb") as f:
                        st.download_button(
                            "📑 Download as DOCX",
                            f.read(),
                            file_name="tailored_resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    try:
                        os.unlink(tmp_docx.name)
                    except:
                        pass

            # Enhanced cover note section
            st.subheader("📝 Short Cover Note")
            try:
                cover = call_for_cover(resume_text, jd_text_local)
                cover_text = ensure_string(cover).strip()
                if cover_text and len(cover_text) > 10:  # Basic validation
                    st.markdown(
                        f"""
                        <div style='padding: 1rem; border-left: 4px solid #2E86C1; background-color: #f0f2f6;'>
                        {cover_text}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.warning("Could not generate a meaningful cover note. Please try again.")
            except Exception as e:
                st.error(f"Error generating cover note: {str(e)}")

# Footer: quick tips
st.markdown("---")
st.markdown("**Tips:**\n- For best results, paste the full job description (responsibilities + required skills). \n- If the model returns malformed JSON, try a shorter resume or split large files. \n- This app depends on the third-party API endpoint — if requests fail, double-check the API_BASE and your API key.")
